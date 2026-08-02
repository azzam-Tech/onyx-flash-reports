import docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = docx.Document()

# Setup styles
style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(12)

# Title
heading = doc.add_heading('ملاحظات على تقرير أعمار التحصيل (أونكس مقابل SREEN) - المندوب 142 لشهر 6', 1)
heading.alignment = WD_ALIGN_PARAGRAPH.RIGHT

# Intro
p = doc.add_paragraph('حياك الله أستاذي،\nبخصوص الفروقات اللي لاحظناها بين تقرير أعمار التحصيل من أونكس والتقرير التحليلي الجديد اللي طلعناه من نظام SREEN، سويت جرد وتدقيق محاسبي عميل بعميل، وطلعت لك بالخلاصة اللي توضح وين الخلل بالضبط وأسباب هذي الفروقات.')
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

doc.add_heading('1. الفارق الإجمالي (وين راحت الفلوس في أونكس؟)', level=2).alignment = WD_ALIGN_PARAGRAPH.RIGHT
p = doc.add_paragraph('الفرق الكلي بين التقريرين هو 357,851.27 ريال مفقودة في إجمالي المندوب بتقرير أونكس. هذي الفلوس ما ضاعت، لكن تقرير أونكس ما حسبها للأسباب التالية:')
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

p = doc.add_paragraph()
p.add_run('أولاً: المبيعات النقدية (238,921.84 ريال): ').bold = True
p.add_run('أونكس متجاهلها تماماً وما حسبها كتحصيل للمندوب مع أنها مدفوعة كاش كجزء من أداء المندوب!')
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

p = doc.add_paragraph()
p.add_run('ثانياً: قيود الشبكة والتسويات (~78,186 ريال): ').bold = True
p.add_run('أونكس يقيس فقط على (سندات القبض النقدية) ويسحب على أي تحصيل يتم عن طريق القيود (مثل سداد الشبكة). هذا خلى عملاء يختفون تماماً من التقرير!')
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

p = doc.add_paragraph()
p.add_run('ثالثاً: مبالغ أسقطها النظام برمجياً (~40,743 ريال): ').bold = True
p.add_run('النظام حق أونكس فيه عيب برمجي غريب، إذا العميل دفع دفعة مقدمة أو سدد دين قديم، بس ما سحب بضاعة جديدة في نفس الشهر (0-30)، أونكس يتبخر المبلغ حقه وما يحطه في التقرير!')
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

doc.add_heading('2. عملاء انظلم فيهم المندوب (بالأرقام)', level=2).alignment = WD_ALIGN_PARAGRAPH.RIGHT
p = doc.add_paragraph('عشان نكون دقيقين، حصرت لك 8 عملاء انحذفت تحصيلاتهم كلياً أو جزئياً من أونكس ومجموع مبالغهم (118,929.43 ريال). وتقدر تراجع كشف حساب أي واحد فيهم بتلقى مبالغه مسجلة بس تقرير أونكس رماها:')
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

table = doc.add_table(rows=1, cols=4)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[3].text = 'اسم ورقم العميل'
hdr_cells[2].text = 'تحصيل SREEN (الصحيح)'
hdr_cells[1].text = 'تحصيل أونكس (الناقص)'
hdr_cells[0].text = 'الفارق (المبلغ المفقود)'

data = [
    ('العميل 1483', '372,141.50', '316,180.50', '55,961.00'),
    ('العميل 2241', '17,310.00', '0.00 (مختفي)', '17,310.00'),
    ('العميل 1988', '97,972.01', '80,703.01', '17,269.00'),
    ('العميل 1043', '7,544.00', '0.00 (مختفي)', '7,544.00'),
    ('العميل 1820', '18,470.00', '12,080.57', '6,389.43'),
    ('العميل 1975', '5,000.00', '0.00 (مختفي)', '5,000.00'),
    ('العميل 1039', '29,042.00', '24,086.00', '4,956.00'),
    ('العميل 1044', '4,500.00', '0.00 (مختفي)', '4,500.00'),
]

for item in data:
    row_cells = table.add_row().cells
    row_cells[3].text = item[0]
    row_cells[2].text = item[1]
    row_cells[1].text = item[2]
    row_cells[0].text = item[3]

doc.add_heading('3. مشكلة توزيع الأعمار (ليش الأرقام تختلف في الخانات؟)', level=2).alignment = WD_ALIGN_PARAGRAPH.RIGHT
p = doc.add_paragraph()
p.add_run('النظام المحاسبي المتعارف عليه يعتمد على مبدأ FIFO (اللي يجي أول يتسدد أول). يعني المفروض نمسح أقدم دين بالكامل وما نترك هللات. وهذا اللي سويناه في ').bold = False
p.add_run('تقرير SREEN').bold = True
p.add_run('، التقرير يجمع كل ما دفعه العميل ويمسح الفواتير الأقدم فالأقدم آلياً دون ما يرمي ولا قرش من حق المندوب أو الشركة.').bold = False
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

p = doc.add_paragraph('أما في أونكس، التوزيع يعتمد على ربط السند بالفاتورة يدوياً، وإذا نسي المحاسب يربط، أونكس يحوس التوزيع! تلقاه يخلي هللات (مثل 0.10 ريال للعميل 1059) في الدين القديم ويحط الباقي بالجديد، أو يقسم المبلغ عشوائياً (مثل ما صار مع العميل 2049 والعميل 1046).')
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

doc.add_heading('الخلاصة', level=2).alignment = WD_ALIGN_PARAGRAPH.RIGHT
p = doc.add_paragraph('طال عمرك، تقرير أونكس لأعمار التحصيل مليان ثغرات وما يعطينا الحسبة الحقيقية لنقدية المندوب لأنه يعتمد على المطابقة اليدوية ويتجاهل قيود الشبكة والمبيعات النقدية. عشان كذا، تقرير SREEN هو الأدق محاسبياً والمحصن ضد الأخطاء البشرية وهو اللي المفروض نعتمد عليه.')
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

# Align all cells right to left if possible, or just leave as default since Word handles it.
for row in table.rows:
    for cell in row.cells:
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

doc.save('C:/Users/amarn/OneDrive/Desktop/dbOnyxOnAntigravity/SREEN_vs_Onyx_Report.docx')
print('Docx created!')
